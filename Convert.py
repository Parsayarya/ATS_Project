import os
from win32com import client
from pytesseract import pytesseract
from docx import Document
import fitz  # PyMuPDF
from PIL import Image
import pytesseract

def convert_pdf_to_docx(folder_path):
    """
    Convert all PDF files in the specified folder to DOCX format using OCR.

    This function iterates over all files in the specified folder, and for each file with a .pdf extension,
    it performs OCR (Optical Character Recognition) to convert it to DOCX format. The DOCX file is saved
    in the same folder with the same base name as the original PDF file.

    :param str folder_path: The path to the folder containing the PDF files to convert.
    """

    # Path to the Tesseract executable
    pytesseract.tesseract_cmd = r"D:\python\ATS\DownloadCorpusBuilder\Tes\tesseract.exe"

    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf"):
            # Path to the PDF file
            pdf_path = os.path.join(folder_path, filename)

            # Create a new Document
            doc = Document()

            # Open the PDF file
            pdf = fitz.open(pdf_path)

            # Loop through each page in the PDF
            for page_num in range(len(pdf)):
                # Get the page
                page = pdf[page_num]

                # Get the image of the page
                image = page.get_pixmap()

                # Get the image as bytes
                imgbytes = image.samples

                # Create a PIL Image object from the bytes data
                image = Image.frombytes("RGB", [image.width, image.height], imgbytes)

                # Use pytesseract to do OCR on the image
                text = pytesseract.image_to_string(image)

                # Add the text to the docx document
                doc.add_paragraph(text)

            # Determine the output file name and save the docx document
            output_filename = os.path.splitext(filename)[0] + ".docx"
            output_filepath = os.path.join(folder_path, output_filename)
            doc.save(output_filepath)


def doc_to_docx(doc_path, docx_path):
    try:
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(docx_path, 16)  # 16 represents the wdFormatDOCX format
        doc.Close()
        word.Quit()
        return True
    except:
        return False


# def doc_to_docx(folder_path):
#     """
#     Convert all .doc files in the specified folder to .docx format using Microsoft Word.
#
#     This function iterates over all files in the specified folder, and for each file with a .doc extension,
#     it opens the file using Microsoft Word, saves it in .docx format, and then closes the Word application.
#
#     :param str folder_path: The path to the folder containing the .doc files to convert.
#     :return: A boolean indicating whether the conversion was successful for all files.
#     :rtype: bool
#     """
#     try:
#         word = client.Dispatch('Word.Application')
#
#         for filename in os.listdir(folder_path):
#             if filename.endswith(".doc"):
#                 # Construct full .doc and .docx file paths
#                 doc_path = os.path.join(folder_path, filename)
#                 docx_path = os.path.join(folder_path, os.path.splitext(filename)[0] + ".docx")
#
#                 # Open the .doc file, save it as .docx, and close it
#                 doc = word.Documents.Open(doc_path)
#                 doc.SaveAs(docx_path, 16)  # 16 represents the wdFormatDOCX format
#                 doc.Close()
#
#         word.Quit()
#         return True
#     except Exception as e:
#         print(f"An error occurred: {e}")
#         return False


# def convert_files_to_docx(folder_path):
#     """
#     Convert all .doc and .pdf files in the specified folder to .docx format.
#
#     This function iterates over all files in the specified folder, converting files with a .doc extension using Microsoft Word,
#     and files with a .pdf extension using OCR (Optical Character Recognition).
#
#     :param str folder_path: The path to the folder containing the .doc and .pdf files to convert.
#     :return: A boolean indicating whether the conversion was successful for all files.
#     :rtype: bool
#     """
#
#     # Path to the Tesseract executable
#     pytesseract.tesseract_cmd = r"D:\python\ATS\DownloadCorpusBuilder\Tes\tesseract.exe"
#
#     try:
#         # DOC to DOCX conversion
#         word = client.Dispatch('Word.Application')
#
#         for filename in os.listdir(folder_path):
#             if filename.endswith(".doc"):
#                 # Construct full .doc and .docx file paths
#                 doc_path = os.path.join(folder_path, filename)
#                 docx_path = os.path.join(folder_path, os.path.splitext(filename)[0] + ".docx")
#
#                 # Open the .doc file, save it as .docx, and close it
#                 doc = word.Documents.Open(doc_path)
#                 doc.SaveAs(docx_path, 16)  # 16 represents the wdFormatDOCX format
#                 doc.Close()
#
#         word.Quit()
#     except Exception as e:
#         print(f"An error occurred during .doc to .docx conversion: {e}")
#         return False
#
#     try:
#         # PDF to DOCX conversion
#         for filename in os.listdir(folder_path):
#             if filename.endswith(".pdf"):
#                 # Path to the PDF file
#                 pdf_path = os.path.join(folder_path, filename)
#
#                 # Create a new Document
#                 doc = Document()
#
#                 # Open the PDF file
#                 pdf = fitz.open(pdf_path)
#
#                 # Loop through each page in the PDF
#                 for page_num in range(len(pdf)):
#                     # Get the page
#                     page = pdf[page_num]
#
#                     # Get the image of the page
#                     image = page.get_pixmap()
#
#                     # Get the image as bytes
#                     imgbytes = image.samples
#
#                     # Create a PIL Image object from the bytes data
#                     image = Image.frombytes("RGB", [image.width, image.height], imgbytes)
#
#                     # Use pytesseract to do OCR on the image
#                     text = pytesseract.image_to_string(image)
#
#                     # Add the text to the docx document
#                     doc.add_paragraph(text)
#
#                 # Determine the output file name and save the docx document
#                 output_filename = os.path.splitext(filename)[0] + ".docx"
#                 output_filepath = os.path.join(folder_path, output_filename)
#                 doc.save(output_filepath)
#
#         return True
#     except Exception as e:
#         print(f"An error occurred during .pdf to .docx conversion: {e}")
#         return False

# from win32com import client
# import os
# from docx import Document
# import fitz
# from PIL import Image



def convert_files_to_docx(folder_path):
    """
    Convert all .doc and .pdf files in the specified folder to .docx format.

    This function iterates over all files in the specified folder, converting files with a .doc extension using Microsoft Word,
    and files with a .pdf extension using OCR (Optical Character Recognition).

    :param str folder_path: The path to the folder containing the .doc and .pdf files to convert.
    :return: A boolean indicating whether the conversion was successful for all files.
    :rtype: bool
    """

    # Path to the Tesseract executable
    pytesseract.pytesseract.tesseract_cmd = r"D:\python\ATS\DownloadCorpusBuilder\Tes\tesseract.exe"

    try:
        # DOC to DOCX conversion
        word = client.Dispatch('Word.Application')
        word.DisplayAlerts = False  # Suppress modal dialogs

        for filename in os.listdir(folder_path):
            if filename.endswith(".doc"):
                # Construct full .doc and .docx file paths
                doc_path = os.path.join(folder_path, filename)
                docx_path = os.path.join(folder_path, os.path.splitext(filename)[0] + ".docx")

                # Open the .doc file, save it to apply repairs, save it as .docx, and close it
                doc = word.Documents.Open(doc_path)
                doc.Save()  # Save once to apply repairs
                doc.SaveAs(docx_path, 16)  # 16 represents the wdFormatDOCX format
                doc.Close()

        word.Quit()
    except Exception as e:
        print(f"An error occurred during .doc to .docx conversion: {e}")
        return False

    try:
        # PDF to DOCX conversion
        for filename in os.listdir(folder_path):
            if filename.endswith(".pdf"):
                # Path to the PDF file
                pdf_path = os.path.join(folder_path, filename)

                # Create a new Document
                doc = Document()

                # Open the PDF file
                pdf = fitz.open(pdf_path)

                # Loop through each page in the PDF
                for page_num in range(len(pdf)):
                    # Get the page
                    page = pdf[page_num]

                    # Get the image of the page
                    image = page.get_pixmap()

                    # Get the image as bytes
                    imgbytes = image.samples

                    # Create a PIL Image object from the bytes data
                    image = Image.frombytes("RGB", [image.width, image.height], imgbytes)

                    # Use pytesseract to do OCR on the image
                    text = pytesseract.image_to_string(image)

                    # Add the text to the docx document
                    doc.add_paragraph(text)

                # Determine the output file name and save the docx document
                output_filename = os.path.splitext(filename)[0] + ".docx"
                output_filepath = os.path.join(folder_path, output_filename)
                doc.save(output_filepath)

        return True
    except Exception as e:
        print(f"An error occurred during .pdf to .docx conversion: {e}")
        return False


convert_files_to_docx(r"D:\python\ATS\DownloadCorpusBuilder\files\docxOutputs\SWPN")
